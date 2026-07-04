"""Mine HIS maturity data from the WHO AFRO HISFA report (HISFA_Report-Final_13022026.docx).

Extracts governance sub-domain scores (Table 2) and data-generation capability
items (Table 4), computes domain means and an overall HIS maturity index, and
assigns the report's maturity bands. Output: an analysis-ready country panel.
"""
from __future__ import annotations

import sys
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
from warehouse import reference  # noqa: E402

REPORT = Path("/mnt/user-data/uploads/HISFA_Report-Final_13022026.docx")
OUT = ROOT / "data" / "processed" / "his"

GOV_COLS = ["Organizational Structure", "Strategic Planning", "Policy Development",
            "Partner Coordination", "HIS Financing", "Supervision National", "Supervision Subnational"]
DATAGEN_HEADER_HINT = "Reporting & Recording"


def _int(x):
    x = x.strip().replace("%", "")
    try:
        return float(x)
    except ValueError:
        return np.nan


def _dedup(cells):
    out = []
    for c in cells:
        c = c.strip()
        if not out or out[-1] != c:
            out.append(c)
    return out


def parse_governance(doc) -> pd.DataFrame:
    rows = []
    for tab in doc.tables:
        head = _dedup([c.text for c in tab.rows[0].cells])
        if head and head[0] == "Country" and "Organizational Structure" in " ".join(head):
            for r in tab.rows[1:]:
                cells = _dedup([c.text for c in r.cells])
                country = canon(cells[0])
                if country in ("Country", "AFRO Average") or "No Data" in " ".join(cells):
                    continue
                vals = [_int(x) for x in cells[1:] if x.strip() != ""]
                sub = vals[:7]
                gov_mean = float(np.nanmean(sub)) if sub else np.nan
                rec = {"country": country, "gov_score": round(gov_mean, 1)}
                for i, cn in enumerate(GOV_COLS):
                    rec[cn] = sub[i] if i < len(sub) else np.nan
                rows.append(rec)
            break
    return pd.DataFrame(rows)


def parse_datagen(doc) -> pd.DataFrame:
    rows = []
    for tab in doc.tables:
        head = _dedup([c.text for c in tab.rows[0].cells])
        if head and head[0] == "Country" and DATAGEN_HEADER_HINT in " ".join(head):
            for r in tab.rows[1:]:
                cells = _dedup([c.text for c in r.cells])
                country = canon(cells[0])
                if country in ("Country", "AFRO"):
                    continue
                vals = [_int(x) for x in cells[1:] if x.strip() not in ("", "-")]
                if not vals:
                    continue
                rows.append({"country": country, "datagen_score": round(float(np.nanmean(vals)), 1)})
            break
    return pd.DataFrame(rows)


NAME_FIX = {
    "The Gambia": "Gambia", "UR Tanzania": "United Republic of Tanzania",
    "Guinea Bissau": "Guinea-Bissau", "Cote d'Ivoire": "Côte d'Ivoire",
}


def canon(country: str) -> str:
    """Canonicalise spelling variants across the two report tables."""
    c = country.strip()
    aliases = {
        "The Gambia": "Gambia", "Guinea Bissau": "Guinea-Bissau",
        "UR Tanzania": "United Republic of Tanzania",
    }
    return aliases.get(c, c)


def _iso(country):
    c = NAME_FIX.get(country, country)
    try:
        return reference.iso3_of(c)
    except Exception:
        # loose match on master list
        master = reference.countries()
        for _, r in master.iterrows():
            if c.lower()[:6] in r["country"].lower() or r["country"].lower()[:6] in c.lower():
                return r["iso3"]
        return None


def band(score):
    if np.isnan(score):
        return "no data"
    return ("nascent" if score <= 20 else "emerging" if score <= 40 else
            "early" if score <= 60 else "advanced" if score <= 80 else "optimized")


def build() -> pd.DataFrame:
    doc = Document(REPORT)
    gov = parse_governance(doc)
    dg = parse_datagen(doc)
    df = gov.merge(dg, on="country", how="outer")
    df["iso3"] = df["country"].map(_iso)
    # overall HIS maturity index = mean of available domain scores (governance, data generation)
    df["his_maturity_index"] = df[["gov_score", "datagen_score"]].mean(axis=1).round(1)
    df["maturity_band"] = df["his_maturity_index"].map(band)
    df["year"] = 2026
    df["source"] = "WHO AFRO HISFA Report 2026 (Tables 2 & 4)"
    return df.sort_values("country")


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    out = build()
    out.to_csv(OUT / "afro_his_maturity.csv", index=False)
    print(f"wrote {len(out)} countries -> {OUT/'afro_his_maturity.csv'}")
    print(out[["country", "gov_score", "datagen_score", "his_maturity_index", "maturity_band"]]
          .to_string(index=False))
    print("\nBand distribution:")
    print(out["maturity_band"].value_counts().to_string())
